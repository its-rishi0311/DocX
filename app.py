from docx import Document

# Create a new document in memory
doc = Document()

# Add sample content to the document
doc.add_heading('DocX Project', level=0)
doc.add_paragraph('Document generated successfully.')

# Save the newly created document
doc.save('output.docx')

# Read and print paragraphs
print("--- Paragraphs ---")
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        print(paragraph.text)

# Read and print table data
print("\n--- Tables ---")
for table in doc.tables:
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        print(row_data)
